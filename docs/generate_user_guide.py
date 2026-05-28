#!/usr/bin/env python3
"""Generate Facturo user guide (DOCX) — IEC/IEEE 26514 aligned."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).parent / "Guide-utilisateur-Facturo.docx"
APP_VERSION = "1.0"
GUIDE_DATE = date(2026, 5, 27).strftime("%d/%m/%Y")
NAVY = RGBColor(0x14, 0x21, 0x3D)
ORANGE = RGBColor(0xFC, 0xA3, 0x11)


def set_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def style_paragraph(p, size=11, bold=False, color=None, space_after=6, align=None):
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align:
        p.alignment = align
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = NAVY
    return h


def add_body(doc, text, bold=False):
    p = doc.add_paragraph(text)
    style_paragraph(p, bold=bold)
    return p


def add_step(doc, number, text):
    p = doc.add_paragraph(f"{number}. {text}", style="List Number")
    style_paragraph(p, space_after=4)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    style_paragraph(p, space_after=4)
    return p


def add_screenshot_box(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"[CAPTURE D'ÉCRAN — {caption}]")
    run.bold = True
    run.italic = True
    run.font.color.rgb = ORANGE
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "FCA311")
        border.append(el)
    p._p.get_or_add_pPr().append(border)


def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{title} — ")
    run.bold = True
    run.font.color.rgb = NAVY
    p.add_run(text)
    style_paragraph(p, space_after=8)


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    hint = doc.add_paragraph(
        "Actualisez la table des matières dans Word : clic droit → Mettre à jour le champ."
    )
    style_paragraph(hint, size=9, color=RGBColor(0x66, 0x66, 0x66))

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run2 = p.add_run()
    run2._r.append(fld_end)


def build_cover(doc):
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title.add_run("FACTURO")
    t_run.bold = True
    t_run.font.size = Pt(36)
    t_run.font.color.rgb = NAVY
    t_run.font.name = "Calibri"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = sub.add_run("Guide utilisateur")
    s_run.font.size = Pt(22)
    s_run.font.color.rgb = ORANGE
    s_run.font.name = "Calibri"

    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tag.add_run("Facturation moderne pour freelances et TPE")
    tag_run.font.size = Pt(14)
    tag_run.font.name = "Calibri"

    for _ in range(4):
        doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in (
        f"Version du guide : {APP_VERSION}",
        f"Version de l'application : {APP_VERSION}",
        f"Date de publication : {GUIDE_DATE}",
        "Public visé : utilisateurs finaux (freelances, TPE, équipes commerciales)",
        "Langue : français",
    ):
        r = meta.add_run(line + "\n")
        r.font.size = Pt(11)
        r.font.name = "Calibri"

    add_page_break(doc)


def build_intro(doc):
    add_heading(doc, "1. Introduction et objectifs du guide", 1)
    add_body(
        doc,
        "Ce guide utilisateur décrit l'utilisation de Facturo, application web de facturation "
        "et de gestion commerciale destinée aux freelances, TPE et petites équipes B2B, "
        "notamment en Afrique de l'Ouest.",
    )
    add_body(
        doc,
        "Facturo permet de créer des devis et des factures conformes, de gérer une base clients, "
        "de suivre les paiements et les impayés, et de piloter votre chiffre d'affaires via un "
        "tableau de bord et des rapports.",
    )

    add_heading(doc, "1.1 Objectifs du guide", 2)
    objectives = [
        "Présenter les prérequis nécessaires à l'utilisation de Facturo.",
        "Décrire les procédures d'inscription, de connexion et de récupération de compte.",
        "Expliquer l'organisation de l'interface et la navigation entre les modules.",
        "Guider pas à pas la réalisation des tâches courantes : clients, devis, factures, paiements, paramètres et abonnement.",
        "Répondre aux questions fréquentes et proposer des solutions en cas de difficulté.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    add_heading(doc, "1.2 Conventions documentaires", 2)
    add_body(doc, "Ce document respecte les principes de la norme IEC/IEEE 26514 :")
    add_bullet(doc, "Les titres sont numérotés (1, 1.1, 1.2…).")
    add_bullet(doc, "Les procédures utilisent des étapes numérotées commençant par un verbe d'action.")
    add_bullet(doc, "Les encadrés [CAPTURE D'ÉCRAN] indiquent les emplacements recommandés pour les illustrations.")
    add_bullet(doc, "Les notes importantes sont signalées par le préfixe « Remarque » ou « Attention ».")

    add_heading(doc, "1.3 Terminologie", 2)
    terms = [
        ("Devis", "Document commercial proposant un prix avant exécution d'une prestation."),
        ("Facture", "Document comptable demandant le règlement d'une prestation."),
        ("Avoir", "Document rectificatif annulant totalement ou partiellement une facture."),
        ("Encours", "Montant total des factures émises mais non encore encaissées."),
        ("Taux de recouvrement", "Pourcentage du chiffre d'affaires facturé effectivement encaissé."),
        ("Plan Gratuit / Pro / Entreprise", "Niveaux d'abonnement définissant les quotas et fonctionnalités."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Terme"
    hdr[1].text = "Définition"
    for term, definition in terms:
        row = table.add_row().cells
        row[0].text = term
        row[1].text = definition

    add_page_break(doc)


def build_prerequisites(doc):
    add_heading(doc, "2. Prérequis et configuration système", 1)

    add_heading(doc, "2.1 Prérequis utilisateur", 2)
    add_bullet(doc, "Disposer d'une adresse e-mail valide et accessible.")
    add_bullet(doc, "Disposer d'une connexion Internet stable.")
    add_bullet(doc, "Connaître les informations légales de votre entreprise (raison sociale, NIF, coordonnées bancaires).")

    add_heading(doc, "2.2 Configuration système recommandée", 2)
    add_body(doc, "Facturo est une application web accessible depuis un navigateur. Aucune installation locale n'est requise.")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Élément"
    table.rows[0].cells[1].text = "Exigence"
    rows = [
        ("Navigateurs supportés", "Google Chrome, Mozilla Firefox, Microsoft Edge ou Safari (versions récentes)"),
        ("Résolution d'écran", "1280 × 720 pixels minimum (1920 × 1080 recommandé)"),
        ("Connexion", "Accès Internet haut débit"),
        ("Périphériques", "Clavier et souris ; écran tactile compatible"),
        ("JavaScript", "Doit être activé dans le navigateur"),
        ("Cookies", "Doivent être autorisés pour la session utilisateur"),
    ]
    for a, b in rows:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b

    add_heading(doc, "2.3 Accès à l'application", 2)
    add_step(doc, 1, "Ouvrez votre navigateur web.")
    add_step(doc, 2, "Saisissez l'adresse https://saas-facturation.vercel.app dans la barre d'adresse.")
    add_step(doc, 3, "Appuyez sur Entrée pour afficher la page d'accueil Facturo.")
    add_screenshot_box(doc, "Page d'accueil Facturo — vue d'ensemble")

    add_note_box(
        doc,
        "Remarque",
        "Facturo fonctionne entièrement en ligne. Vos données sont synchronisées avec le serveur "
        "à chaque action ; une connexion Internet est indispensable.",
    )
    add_page_break(doc)


def build_auth(doc):
    add_heading(doc, "3. Inscription et connexion", 1)

    add_heading(doc, "3.1 Créer un compte", 2)
    add_step(doc, 1, "Accédez à la page d'accueil Facturo.")
    add_step(doc, 2, "Cliquez sur « Commencer gratuitement » ou « S'inscrire ».")
    add_step(doc, 3, "Renseignez votre nom complet.")
    add_step(doc, 4, "Saisissez votre adresse e-mail professionnelle.")
    add_step(doc, 5, "Choisissez un mot de passe respectant les critères de sécurité :")
    add_bullet(doc, "8 caractères minimum")
    add_bullet(doc, "Au moins une lettre")
    add_bullet(doc, "Au moins un chiffre")
    add_bullet(doc, "Au moins un symbole")
    add_step(doc, 6, "Confirmez votre mot de passe.")
    add_step(doc, 7, "Cliquez sur « Créer mon compte ».")
    add_step(doc, 8, "Consultez votre boîte e-mail et cliquez sur le lien de vérification reçu.")
    add_screenshot_box(doc, "Formulaire d'inscription")

    add_note_box(
        doc,
        "Remarque",
        "À l'inscription, votre compte est automatiquement rattaché au plan Gratuit. "
        "Vous pouvez passer à l'offre Pro à tout moment depuis la section Abonnement.",
    )

    add_heading(doc, "3.2 Se connecter", 2)
    add_step(doc, 1, "Cliquez sur « Se connecter » depuis la page d'accueil.")
    add_step(doc, 2, "Saisissez votre adresse e-mail.")
    add_step(doc, 3, "Saisissez votre mot de passe.")
    add_step(doc, 4, "Cliquez sur « Connexion ».")
    add_step(doc, 5, "Accédez au tableau de bord (/app) une fois authentifié.")
    add_screenshot_box(doc, "Page de connexion")

    add_heading(doc, "3.3 Mot de passe oublié", 2)
    add_step(doc, 1, "Cliquez sur « Mot de passe oublié ? » depuis la page de connexion.")
    add_step(doc, 2, "Saisissez l'adresse e-mail associée à votre compte.")
    add_step(doc, 3, "Cliquez sur « Envoyer le lien de réinitialisation ».")
    add_step(doc, 4, "Ouvrez l'e-mail reçu et cliquez sur le lien de réinitialisation.")
    add_step(doc, 5, "Saisissez votre nouveau mot de passe et confirmez-le.")
    add_step(doc, 6, "Cliquez sur « Réinitialiser le mot de passe ».")
    add_step(doc, 7, "Connectez-vous avec vos nouveaux identifiants.")

    add_heading(doc, "3.4 Se déconnecter", 2)
    add_step(doc, 1, "Cliquez sur votre avatar ou votre nom en haut à droite de l'interface.")
    add_step(doc, 2, "Sélectionnez « Déconnexion » dans le menu déroulant.")
    add_step(doc, 3, "Vérifiez que vous êtes redirigé vers la page d'accueil publique.")

    add_heading(doc, "3.5 Sécurité du compte", 2)
    add_bullet(doc, "Ne partagez jamais votre mot de passe.")
    add_bullet(doc, "Utilisez un mot de passe unique pour Facturo.")
    add_bullet(doc, "Déconnectez-vous après utilisation sur un poste partagé.")
    add_bullet(doc, "Activez la protection des montants sur le tableau de bord en environnement partagé (voir section 7.1).")
    add_page_break(doc)


def build_interface(doc):
    add_heading(doc, "4. Présentation de l'interface", 1)

    add_heading(doc, "4.1 Vue d'ensemble", 2)
    add_body(
        doc,
        "L'application authentifiée est organisée autour d'une barre latérale (sidebar), "
        "d'un en-tête (header) et d'une zone de contenu principal.",
    )
    add_screenshot_box(doc, "Interface principale — sidebar, header et zone de contenu")

    add_heading(doc, "4.2 Barre latérale (navigation)", 2)
    add_body(doc, "La barre latérale donne accès aux modules principaux :")
    modules = [
        ("Tableau de bord", "Vue synthétique de votre activité commerciale."),
        ("Factures", "Création, suivi et gestion des factures."),
        ("Devis", "Création et suivi des propositions commerciales."),
        ("Clients", "Gestion de votre carnet d'adresses clients."),
        ("Rapports", "Analyses avancées et exports (offre Pro)."),
        ("Paramètres", "Configuration de l'entreprise et des documents."),
    ]
    for name, desc in modules:
        add_bullet(doc, f"{name} — {desc}")

    add_step(doc, 1, "Cliquez sur un élément du menu pour accéder au module correspondant.")
    add_step(doc, 2, "Cliquez sur l'icône de réduction pour replier la barre latérale et gagner de l'espace.")
    add_note_box(doc, "Remarque", "L'état replié/déplié de la sidebar est mémorisé dans votre navigateur.")

    add_heading(doc, "4.3 En-tête", 2)
    add_bullet(doc, "Titre de la page courante.")
    add_bullet(doc, "Badge indiquant votre plan d'abonnement (Gratuit, Pro, Entreprise).")
    add_bullet(doc, "Menu profil : Mon profil, Abonnement, Déconnexion.")

    add_heading(doc, "4.4 Éléments d'interface récurrents", 2)
    add_bullet(doc, "Boutons d'action principaux (« Nouveau devis », « Nouvelle facture », etc.).")
    add_bullet(doc, "Barres de recherche et filtres par statut.")
    add_bullet(doc, "Modales de création, édition et confirmation.")
    add_bullet(doc, "Indicateurs de chargement (squelettes) pendant le chargement des données.")
    add_bullet(doc, "Messages de confirmation avant les actions irréversibles.")

    add_heading(doc, "4.5 Langue et fuseau horaire", 2)
    add_step(doc, 1, "Accédez à Paramètres > Préférences.")
    add_step(doc, 2, "Sélectionnez la langue de l'interface (Français ou Anglais).")
    add_step(doc, 3, "Choisissez votre fuseau horaire (Abidjan, Lagos, Casablanca, Paris, UTC).")
    add_step(doc, 4, "Cliquez sur « Enregistrer » pour appliquer les modifications.")
    add_page_break(doc)


def build_features(doc):
    add_heading(doc, "5. Fonctionnalités principales", 1)

    # 5.1 Dashboard
    add_heading(doc, "5.1 Tableau de bord", 2)
    add_body(doc, "Le tableau de bord offre une vue d'ensemble de votre activité commerciale.")
    add_screenshot_box(doc, "Tableau de bord — indicateurs et graphiques")

    add_heading(doc, "5.1.1 Consulter les indicateurs clés", 3)
    add_body(doc, "Six indicateurs (KPI) sont affichés :")
    kpis = [
        "Chiffre d'affaires encaissé",
        "Encours (montants non encore perçus)",
        "Taux de recouvrement",
        "Montant moyen par facture",
        "Nombre de clients actifs",
        "Nombre de factures en retard",
    ]
    for k in kpis:
        add_bullet(doc, k)

    add_heading(doc, "5.1.2 Déverrouiller l'affichage des montants", 3)
    add_step(doc, 1, "Accédez au tableau de bord.")
    add_step(doc, 2, "Cliquez sur l'icône de cadenas si les montants sont masqués (******).")
    add_step(doc, 3, "Saisissez votre mot de passe de connexion.")
    add_step(doc, 4, "Cliquez sur « Valider » pour afficher les montants.")

    add_heading(doc, "5.1.3 Exporter les données du tableau de bord (Pro)", 3)
    add_step(doc, 1, "Vérifiez que votre compte est sur l'offre Pro.")
    add_step(doc, 2, "Cliquez sur « Exporter CSV » depuis le tableau de bord.")
    add_step(doc, 3, "Enregistrez le fichier sur votre poste.")

    # 5.2 Clients
    add_heading(doc, "5.2 Gestion des clients", 2)
    add_step(doc, 1, "Cliquez sur « Clients » dans la barre latérale.")
    add_step(doc, 2, "Cliquez sur « Nouveau client ».")
    add_step(doc, 3, "Renseignez les champs : prénom, nom, e-mail, téléphone, société, adresse, NIF, notes.")
    add_step(doc, 4, "Cliquez sur « Enregistrer ».")
    add_screenshot_box(doc, "Formulaire de création client")

    add_heading(doc, "5.2.1 Rechercher et filtrer les clients", 3)
    add_step(doc, 1, "Saisissez un terme dans la barre de recherche (nom, société, e-mail).")
    add_step(doc, 2, "Appliquez les filtres disponibles si nécessaire.")
    add_step(doc, 3, "Cliquez sur un client pour ouvrir sa fiche détaillée.")

    add_heading(doc, "5.2.2 Consulter l'historique d'un client", 3)
    add_step(doc, 1, "Ouvrez la fiche du client concerné.")
    add_step(doc, 2, "Consultez le chiffre d'affaires encaissé.")
    add_step(doc, 3, "Parcourez la liste des devis et factures associés.")

    add_heading(doc, "5.2.3 Importer des clients via CSV (Pro)", 3)
    add_step(doc, 1, "Cliquez sur « Importer » depuis la page Clients.")
    add_step(doc, 2, "Préparez un fichier CSV avec les colonnes requises.")
    add_step(doc, 3, "Collez le contenu CSV ou sélectionnez le fichier.")
    add_step(doc, 4, "Cliquez sur « Importer » pour valider l'importation.")
    add_screenshot_box(doc, "Modal d'import CSV clients")

    # 5.3 Devis
    add_heading(doc, "5.3 Gestion des devis", 2)
    add_step(doc, 1, "Cliquez sur « Devis » dans la barre latérale.")
    add_step(doc, 2, "Cliquez sur « Nouveau devis ».")
    add_step(doc, 3, "Sélectionnez le client destinataire.")
    add_step(doc, 4, "Indiquez la date d'émission et la date de validité.")
    add_step(doc, 5, "Choisissez la devise (XOF, EUR ou USD).")
    add_step(doc, 6, "Ajoutez une ou plusieurs lignes de prestation (description, quantité, prix unitaire, TVA, remise).")
    add_step(doc, 7, "Vérifiez les totaux HT, TVA et TTC calculés automatiquement.")
    add_step(doc, 8, "Ajoutez des notes complémentaires si nécessaire.")
    add_step(doc, 9, "Cliquez sur « Enregistrer » — le devis est créé en statut Brouillon.")
    add_screenshot_box(doc, "Formulaire de création de devis")

    add_heading(doc, "5.3.1 Cycle de vie d'un devis", 3)
    add_body(doc, "Statuts possibles : Brouillon → Envoyé → Accepté / Refusé / Expiré.")
    add_step(doc, 1, "Ouvrez le devis concerné.")
    add_step(doc, 2, "Cliquez sur le bouton de changement de statut.")
    add_step(doc, 3, "Sélectionnez le nouveau statut.")
    add_step(doc, 4, "Confirmez le changement dans la fenêtre de dialogue.")

    add_heading(doc, "5.3.2 Prévisualiser et télécharger un devis en PDF", 3)
    add_step(doc, 1, "Ouvrez le devis concerné.")
    add_step(doc, 2, "Cliquez sur « Aperçu » pour visualiser le document.")
    add_step(doc, 3, "Cliquez sur « Télécharger PDF » pour enregistrer le fichier.")

    add_heading(doc, "5.3.3 Convertir un devis en facture", 3)
    add_note_box(
        doc,
        "Attention",
        "La conversion en facture est une action définitive et irréversible.",
    )
    add_step(doc, 1, "Ouvrez un devis accepté ou validé.")
    add_step(doc, 2, "Cliquez sur « Convertir en facture ».")
    add_step(doc, 3, "Confirmez l'action dans la fenêtre de dialogue.")
    add_step(doc, 4, "Accédez à la section Factures pour consulter la facture générée.")

    # 5.4 Factures
    add_heading(doc, "5.4 Gestion des factures", 2)
    add_heading(doc, "5.4.1 Créer une facture manuellement", 3)
    add_step(doc, 1, "Cliquez sur « Factures » dans la barre latérale.")
    add_step(doc, 2, "Cliquez sur « Nouvelle facture ».")
    add_step(doc, 3, "Sélectionnez le client.")
    add_step(doc, 4, "Indiquez les dates d'émission et d'échéance.")
    add_step(doc, 5, "Choisissez la devise.")
    add_step(doc, 6, "Ajoutez les lignes ou laissez les montants être calculés depuis un devis source.")
    add_step(doc, 7, "Cliquez sur « Enregistrer ».")
    add_screenshot_box(doc, "Liste des factures et formulaire de création")

    add_heading(doc, "5.4.2 Cycle de vie d'une facture", 3)
    add_body(doc, "Statuts possibles : Brouillon → Envoyée → Payée / En retard / Annulée.")
    add_step(doc, 1, "Ouvrez la facture concernée.")
    add_step(doc, 2, "Modifiez le statut via le menu dédié.")
    add_step(doc, 3, "Confirmez le changement.")

    add_note_box(
        doc,
        "Remarque",
        "Les factures dont la date d'échéance est dépassée passent automatiquement en statut « En retard » "
        "(traitement quotidien côté serveur).",
    )

    add_heading(doc, "5.4.3 Enregistrer un paiement", 3)
    add_step(doc, 1, "Ouvrez la facture concernée.")
    add_step(doc, 2, "Cliquez sur « Paiements » ou l'icône associée.")
    add_step(doc, 3, "Saisissez le montant perçu.")
    add_step(doc, 4, "Indiquez la méthode de paiement (virement, espèces, chèque, mobile money, etc.).")
    add_step(doc, 5, "Renseignez la référence et la date du paiement.")
    add_step(doc, 6, "Cliquez sur « Enregistrer ».")
    add_step(doc, 7, "Vérifiez le solde restant ; la facture passe à « Payée » lorsque le solde atteint zéro.")
    add_screenshot_box(doc, "Modal d'enregistrement de paiement")

    add_heading(doc, "5.4.4 Télécharger une facture en PDF", 3)
    add_step(doc, 1, "Ouvrez la facture concernée.")
    add_step(doc, 2, "Cliquez sur « Aperçu ».")
    add_step(doc, 3, "Cliquez sur « Télécharger PDF ».")

    add_heading(doc, "5.4.5 Créer un avoir", 3)
    add_note_box(doc, "Attention", "La création d'un avoir est une action définitive.")
    add_step(doc, 1, "Ouvrez la facture à rectifier.")
    add_step(doc, 2, "Cliquez sur « Créer un avoir ».")
    add_step(doc, 3, "Confirmez l'action.")

    add_heading(doc, "5.4.6 Limite du plan Gratuit", 3)
    add_body(
        doc,
        "Le plan Gratuit autorise la création de 10 factures par mois calendaire. "
        "Au-delà de cette limite, la création est bloquée. Passez à l'offre Pro pour des factures illimitées.",
    )

    # 5.5 Rapports
    add_heading(doc, "5.5 Rapports et analyses", 2)
    add_step(doc, 1, "Cliquez sur « Rapports » dans la barre latérale.")
    add_step(doc, 2, "Sélectionnez la période d'analyse (mois, trimestre ou année).")
    add_step(doc, 3, "Consultez les indicateurs, graphiques et classements.")
    add_screenshot_box(doc, "Page Rapports — graphiques et KPI")

    add_heading(doc, "5.5.1 Fonctionnalités avancées (Pro)", 3)
    add_bullet(doc, "Détail par client")
    add_bullet(doc, "Analyse d'ancienneté des retards (aging)")
    add_bullet(doc, "Export CSV des rapports")

    add_step(doc, 1, "Vérifiez que votre compte est sur l'offre Pro.")
    add_step(doc, 2, "Utilisez les filtres avancés disponibles.")
    add_step(doc, 3, "Cliquez sur « Exporter CSV » pour télécharger les données.")

    # 5.6 Paramètres
    add_heading(doc, "5.6 Paramètres de l'entreprise et des documents", 2)
    add_step(doc, 1, "Cliquez sur « Paramètres » dans la barre latérale.")
    add_step(doc, 2, "Renseignez les informations de votre entreprise : raison sociale, adresse, téléphone, e-mail affiché.")
    add_step(doc, 3, "Saisissez votre NIF ou identifiant fiscal.")
    add_step(doc, 4, "Indiquez vos coordonnées bancaires (banque, IBAN, BIC).")
    add_step(doc, 5, "Rédigez les mentions légales affichées en pied de page des PDF.")
    add_step(doc, 6, "Téléversez votre logo (format image, 2 Mo maximum).")
    add_step(doc, 7, "Personnalisez les couleurs des documents (couleur primaire et accent).")
    add_step(doc, 8, "Cliquez sur « Enregistrer ».")
    add_screenshot_box(doc, "Page Paramètres — entreprise et branding")

    add_heading(doc, "5.6.1 Activer les notifications par e-mail", 3)
    add_step(doc, 1, "Accédez à Paramètres > Préférences.")
    add_step(doc, 2, "Activez l'option « Notifications e-mail ».")
    add_step(doc, 3, "Enregistrez vos modifications.")
    add_note_box(
        doc,
        "Remarque",
        "Avec les notifications activées et un abonnement Pro, Facturo envoie des relances automatiques "
        "hebdomadaires pour les factures en retard (chaque lundi à 8 h, fuseau serveur).",
    )

    # 5.7 Profil
    add_heading(doc, "5.7 Mon profil", 2)
    add_step(doc, 1, "Cliquez sur votre avatar en haut à droite.")
    add_step(doc, 2, "Sélectionnez « Mon profil ».")
    add_step(doc, 3, "Modifiez votre nom ou votre adresse e-mail si nécessaire.")
    add_step(doc, 4, "Pour changer votre mot de passe, saisissez l'ancien mot de passe puis le nouveau.")
    add_step(doc, 5, "Cliquez sur « Enregistrer ».")

    # 5.8 Abonnement
    add_heading(doc, "5.8 Gestion de l'abonnement", 2)
    add_body(doc, "Facturo propose trois offres :")
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Offre"
    table.rows[0].cells[1].text = "Tarif"
    table.rows[0].cells[2].text = "Principales caractéristiques"
    offers = [
        ("Gratuit", "0 F CFA/mois", "10 factures/mois, PDF, suivi client basique"),
        ("Pro", "5 000 F CFA/mois", "Factures illimitées, relances auto, exports CSV, rapports avancés, support prioritaire"),
        ("Entreprise", "Sur devis", "Multi-utilisateurs, SLA 99,5 %, onboarding, intégrations"),
    ]
    for a, b, c in offers:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b
        r[2].text = c

    add_heading(doc, "5.8.1 Passer à l'offre Pro", 3)
    add_step(doc, 1, "Cliquez sur votre avatar > « Abonnement ».")
    add_step(doc, 2, "Sélectionnez l'offre Pro.")
    add_step(doc, 3, "Cliquez sur « Passer à Pro ».")
    add_step(doc, 4, "Suivez les instructions de paiement Stripe (carte bancaire).")
    add_step(doc, 5, "Attendez la confirmation de paiement et vérifiez la mise à jour de votre plan.")
    add_screenshot_box(doc, "Page Abonnement — comparaison des offres")

    add_heading(doc, "5.8.2 Gérer votre abonnement Pro", 3)
    add_step(doc, 1, "Accédez à la page Abonnement.")
    add_step(doc, 2, "Cliquez sur « Gérer mon abonnement » pour ouvrir le portail Stripe.")
    add_step(doc, 3, "Modifiez votre moyen de paiement ou résiliez selon vos besoins.")

    add_page_break(doc)


def build_faq(doc):
    add_heading(doc, "6. Foire aux questions (FAQ)", 1)

    faqs = [
        (
            "6.1 Facturo est-il conforme aux exigences de facturation ?",
            "Oui. Facturo génère des documents structurés avec numérotation chronologique, "
            "mentions légales configurables et calcul automatique des montants HT, TVA et TTC.",
        ),
        (
            "6.2 Puis-je importer mes clients existants ?",
            "Oui, via l'import CSV disponible avec l'offre Pro (section 5.2.3).",
        ),
        (
            "6.3 Le plan Gratuit est-il limité dans le temps ?",
            "Non. Le plan Gratuit est disponible sans limite de durée, mais limité à 10 factures par mois.",
        ),
        (
            "6.4 Quelles devises sont prises en charge ?",
            "Facturo prend en charge le franc CFA (XOF) par défaut, ainsi que l'euro (EUR) et le dollar US (USD).",
        ),
        (
            "6.5 Les relances de paiement sont-elles automatiques ?",
            "Oui, avec l'offre Pro et les notifications e-mail activées. Des relances hebdomadaires sont "
            "envoyées pour les factures en retard.",
        ),
        (
            "6.6 Mes clients peuvent-ils payer en ligne via Facturo ?",
            "Facturo permet d'enregistrer manuellement les paiements reçus. Le paiement en ligne direct "
            "par vos clients n'est pas disponible dans la version actuelle.",
        ),
        (
            "6.7 Puis-je partager mes données avec mon comptable ?",
            "Oui. Exportez vos données en CSV (offre Pro) ou téléchargez les PDF de vos factures "
            "et transmettez-les à votre comptable.",
        ),
        (
            "6.8 Mes données sont-elles sécurisées ?",
            "Oui. Chaque compte est isolé. L'authentification repose sur des tokens sécurisés. "
            "Consultez la politique de confidentialité sur /legal/confidentialite.",
        ),
        (
            "6.9 Puis-je utiliser Facturo sur mobile ?",
            "Oui. L'interface web est responsive et utilisable depuis un navigateur mobile récent.",
        ),
        (
            "6.10 Comment contacter le support ?",
            "Envoyez un e-mail à contact@facturo.app. Les abonnés Pro bénéficient d'un support prioritaire.",
        ),
    ]
    for title, answer in faqs:
        add_heading(doc, title, 2)
        add_body(doc, answer)

    add_page_break(doc)


def build_troubleshooting(doc):
    add_heading(doc, "7. Dépannage", 1)
    add_body(
        doc,
        "Cette section décrit les problèmes les plus fréquents et les actions correctives recommandées.",
    )

    issues = [
        (
            "7.1 Impossible de se connecter",
            [
                "Vérifiez que votre adresse e-mail et votre mot de passe sont corrects.",
                "Utilisez la procédure « Mot de passe oublié » (section 3.3).",
                "Videz le cache de votre navigateur et réessayez.",
                "Essayez un autre navigateur supporté.",
            ],
        ),
        (
            "7.2 Session expirée ou déconnexion automatique",
            [
                "Reconnectez-vous avec vos identifiants.",
                "Vérifiez que les cookies ne sont pas bloqués par votre navigateur.",
                "Évitez d'utiliser la navigation privée si elle efface les cookies à la fermeture.",
            ],
        ),
        (
            "7.3 L'application ne charge pas les données",
            [
                "Vérifiez votre connexion Internet.",
                "Actualisez la page (F5 ou Ctrl+R).",
                "Déconnectez-vous puis reconnectez-vous.",
                "Si le problème persiste, consultez la page de statut ou contactez le support.",
            ],
        ),
        (
            "7.4 Impossible de créer une nouvelle facture",
            [
                "Vérifiez si vous avez atteint la limite de 10 factures/mois (plan Gratuit).",
                "Passez à l'offre Pro pour lever cette limitation (section 5.8.1).",
            ],
        ),
        (
            "7.5 Export CSV ou import clients indisponible",
            [
                "Ces fonctionnalités sont réservées à l'offre Pro.",
                "Vérifiez votre plan actuel dans Mon profil ou Abonnement.",
            ],
        ),
        (
            "7.6 Rapports avancés verrouillés",
            [
                "Les analyses détaillées et l'export CSV des rapports nécessitent l'offre Pro.",
                "Mettez à jour votre abonnement depuis la page Abonnement.",
            ],
        ),
        (
            "7.7 Montants masqués sur le tableau de bord",
            [
                "Cliquez sur l'icône de cadenas.",
                "Saisissez votre mot de passe pour déverrouiller l'affichage (section 5.1.2).",
            ],
        ),
        (
            "7.8 Échec du paiement Stripe lors du passage à Pro",
            [
                "Vérifiez que votre carte bancaire est valide et dispose de fonds suffisants.",
                "Réessayez depuis la page Abonnement.",
                "Contactez contact@facturo.app si le problème persiste.",
            ],
        ),
        (
            "7.9 Le PDF ne s'affiche pas ou est incomplet",
            [
                "Vérifiez que les paramètres de votre entreprise sont renseignés (section 5.6).",
                "Désactivez temporairement les bloqueurs de pop-up.",
                "Essayez un autre navigateur.",
            ],
        ),
        (
            "7.10 E-mail de vérification non reçu",
            [
                "Vérifiez votre dossier spam ou courrier indésirable.",
                "Assurez-vous que l'adresse e-mail saisie est correcte.",
                "Connectez-vous et demandez un nouvel envoi de l'e-mail de vérification.",
            ],
        ),
    ]

    for title, steps in issues:
        add_heading(doc, title, 2)
        for i, step in enumerate(steps, 1):
            add_step(doc, i, step)

    add_page_break(doc)


def build_support(doc):
    add_heading(doc, "8. Contact et support", 1)

    add_heading(doc, "8.1 Coordonnées", 2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Type de demande"
    table.rows[0].cells[1].text = "Contact"
    contacts = [
        ("Support général", "contact@facturo.app"),
        ("Confidentialité et RGPD", "privacy@facturo.app"),
        ("Offre Entreprise", "contact@facturo.app (objet : Offre Entreprise Facturo)"),
        ("Site web", "https://saas-facturation.vercel.app"),
    ]
    for a, b in contacts:
        r = table.add_row().cells
        r[0].text = a
        r[1].text = b

    add_heading(doc, "8.2 Niveaux de support", 2)
    add_bullet(doc, "Plan Gratuit : support communautaire par e-mail.")
    add_bullet(doc, "Plan Pro : support prioritaire par e-mail.")
    add_bullet(doc, "Plan Entreprise : support dédié avec SLA et onboarding personnalisé.")

    add_heading(doc, "8.3 Informations à fournir lors d'une demande", 2)
    add_body(doc, "Pour accélérer le traitement de votre demande, indiquez :")
    add_bullet(doc, "Votre adresse e-mail de connexion.")
    add_bullet(doc, "Une description précise du problème rencontré.")
    add_bullet(doc, "Les étapes déjà tentées (section 7).")
    add_bullet(doc, "Le navigateur et le système d'exploitation utilisés.")
    add_bullet(doc, "Des captures d'écran si possible.")

    add_heading(doc, "8.4 Documents légaux", 2)
    add_bullet(doc, "Mentions légales : /legal/mentions")
    add_bullet(doc, "Politique de confidentialité : /legal/confidentialite")

    add_heading(doc, "8.5 Historique du document", 2)
    history = doc.add_table(rows=1, cols=4)
    history.style = "Table Grid"
    history.rows[0].cells[0].text = "Version"
    history.rows[0].cells[1].text = "Date"
    history.rows[0].cells[2].text = "Auteur"
    history.rows[0].cells[3].text = "Modifications"
    row = history.add_row().cells
    row[0].text = APP_VERSION
    row[1].text = GUIDE_DATE
    row[2].text = "Équipe Facturo"
    row[3].text = "Publication initiale du guide utilisateur"


def main():
    doc = Document()
    set_margins(doc)

    build_cover(doc)

    add_heading(doc, "Table des matières", 1)
    add_toc(doc)
    add_page_break(doc)

    build_intro(doc)
    build_prerequisites(doc)
    build_auth(doc)
    build_interface(doc)
    build_features(doc)
    build_faq(doc)
    build_troubleshooting(doc)
    build_support(doc)

    doc.save(OUTPUT)
    print(f"Guide généré : {OUTPUT}")


if __name__ == "__main__":
    main()
